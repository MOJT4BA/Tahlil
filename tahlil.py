import streamlit as st
import pandas as pd
import numpy as np
from docx import Document

# تعریف فاکتورها و ضرایب
fundamental_factors = [
    "محبوبیت رسانه",
    "محتوای خبری (مثبت/منفی)",
    "نوع و کیفیت نظرات",
    "تحلیل احساسات (Sentiment Analysis)",
    "رویدادهای مرتبط",
    "تحلیل ترندهای اجتماعی",
    "پوشش خبری توسط رسانه‌های معتبر",
    "ارزیابی کمپین‌های تبلیغاتی",
    "تأثیر اخبار خاص",
    "تحلیل رفتار کاربران",
    "تحلیل کیفیت محتوا",
    "واکنش‌های فوری کاربران",
    "ارزیابی منابع خبری",
    "تحلیل تأثیرات اقتصادی",
    "تحلیل پایداری اخبار"
]

technical_factors = [
    "تعداد تکرار هشتگ‌ها (#hashtags)",
    "تعداد لایک‌ها",
    "تعداد دیس‌لایک‌ها",
    "تعداد نظرات",
    "تعداد اشتراک‌گذاری‌ها",
    "تعداد بازدیدها",
    "تعداد ری‌توییت‌ها",
    "تعداد دنبال‌کنندگان جدید",
    "تعداد انتشار مقالات در رسانه‌ها",
    "میزان تعامل (Engagement rate)",
    "تعداد پست‌ها در شبکه‌های اجتماعی",
    "نرخ کلیک (Click-through rate)",
    "تعداد مشاهده ویدئوها",
    "تعداد انتشار پست‌ها در وبلاگ‌ها",
    "تعداد مقالات منتشر شده در مجلات",
    "تعداد دانلودها و استریم‌ها"
]

# رابط کاربری
st.title("تحلیل خبر و تعیین ضرایب")

news_content = st.text_area("متن خبر", "خودروهای برقی می توانند وارد محدوده طرح ترافیک شوند بدون آن که جریمه شوند")

st.header("داده‌های فاندامنتال")
fundamental_data = []
for factor in fundamental_factors:
    weight = st.slider(f"ضریب {factor}", 0.0, 1.0, 0.1)
    score = st.slider(f"نمره {factor}", 0.0, 1.0, 0.1)
    fundamental_data.append((factor, weight, score))

st.header("داده‌های تکنیکال")
technical_data = []
for factor in technical_factors:
    weight = st.slider(f"ضریب {factor}", 0.0, 1.0, 0.1)
    score = st.slider(f"نمره {factor}", 0.0, 1.0, 0.1)
    technical_data.append((factor, weight, score))

# محاسبه نمره کلی
fundamental_score = sum(weight * score for _, weight, score in fundamental_data)
technical_score = sum(weight * score for _, weight, score in technical_data)
overall_score = (fundamental_score + technical_score) / 2
normalized_score = overall_score * 10 / max(fundamental_score, technical_score)

st.subheader("نتایج")
st.write(f"نمره کلی داده‌های فاندامنتال: {fundamental_score:.2f}")
st.write(f"نمره کلی داده‌های تکنیکال: {technical_score:.2f}")
st.write(f"نمره کلی جریان‌سازی: {normalized_score:.2f} از 10")

# ایجاد فایل Word
if st.button("ایجاد فایل Word"):
    doc = Document()
    doc.add_heading('تحلیل خبر: "خودروهای برقی می توانند وارد محدوده طرح ترافیک شوند بدون آن که جریمه شوند"', level=1)

    doc.add_heading('جدول داده‌های فاندامنتال', level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'فاکتور فاندامنتال'
    hdr_cells[1].text = 'ضریب'
    hdr_cells[2].text = 'نمره'
    hdr_cells[3].text = 'علت انتخاب نمره و ضریب'
    for factor, weight, score in fundamental_data:
        row_cells = table.add_row().cells
        row_cells[0].text = factor
        row_cells[1].text = str(weight)
        row_cells[2].text = str(score)
        row_cells[3].text = f'علت نمره: {score}'

    doc.add_paragraph(f'نمره کلی داده‌های فاندامنتال: {fundamental_score:.2f}')

    doc.add_heading('جدول داده‌های تکنیکال', level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'فاکتور تکنیکال'
    hdr_cells[1].text = 'ضریب'
    hdr_cells[2].text = 'نمره'
    hdr_cells[3].text = 'علت انتخاب نمره و ضریب'
    for factor, weight, score in technical_data:
        row_cells = table.add_row().cells
        row_cells[0].text = factor
        row_cells[1].text = str(weight)
        row_cells[2].text = str(score)
        row_cells[3].text = f'علت نمره: {score}'

    doc.add_paragraph(f'نمره کلی داده‌های تکنیکال: {technical_score:.2f}')
    doc.add_paragraph(f'نمره کلی جریان‌سازی: {normalized_score:.2f} از 10')

    doc.save("News_Analysis.docx")
    st.success("فایل Word ایجاد شد!")
    with open("News_Analysis.docx", "rb") as file:
        st.download_button("دانلود فایل Word", file, file_name="News_Analysis.docx")

